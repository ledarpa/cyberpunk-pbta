#!/usr/bin/env python3
"""Ensambla capitulos Markdown en cyberpunk-pbta.docx — tema MS-DOS / 2 columnas.

Jerarquía Word nativa (TOC):
  #     → Heading 1 (Título 1)
  ##    → Heading 2 (Título 2)
  ###   → Heading 3 (Título 3)
  ####  → Heading 4
  texto → Normal (párrafo)
  -     → List Bullet (lista)
  indent+ - → List Bullet 2 / nivel 1 (sublista)
  |…|   → Table Grid + fila cabecera
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SCRIPTS = Path(__file__).resolve().parent
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from word_styles import (  # noqa: E402
    STYLE_H2,
    STYLE_H3,
    STYLE_H4,
    STYLE_MESA,
    STYLE_NORMAL,
    STYLE_QUOTE,
    add_body_paragraph,
    add_bullet,
    add_cover_art,
    add_part_title_ascii,
    add_rule_separator,
    apply_all_section_margins,
    set_paragraph_keep,
    setup_document_styles,
    start_body_layout,
    style_table,
)

CAPITULOS = ROOT / "docs" / "capitulos"
PORTADA = ROOT / "docs" / "assets" / "portada-ascii.txt"
OUTPUT = ROOT / "cyberpunk-pbta.docx"

CHAPTER_FILES = [
    "00-sistema.md",
    "01-crear-un-cyberpunk.md",
    "02-cyberware-reglas-y-economia.md",
    "04-catalogo-cromos.md",
    "05-catalogo-chaperia.md",
    "06-glosario.md",
]


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


_SECOND_HEADER_COL0 = {
    "módulo",
    "modulos",
    "módulos",
    "accesorio",
    "accesorios",
    "subsistema",
    "subsistemas",
}
_SECOND_HEADER_COL1 = {
    "efecto",
    "efectos",
    "reglas",
    "módulos",
    "modulos",
    "accesorios",
    "subsistemas",
}


def split_dual_title_tables(rows: list[list[str]]) -> list[list[list[str]]]:
    """Parte tablas fusionadas Calidad|… + Módulo|Efecto (u homólogas) en 2 tablas."""
    if len(rows) < 4:
        return [rows]
    chunks: list[list[list[str]]] = []
    current: list[list[str]] = [rows[0]]
    for row in rows[1:]:
        c0 = (row[0] if row else "").strip().lower()
        c1 = (row[1] if len(row) > 1 else "").strip().lower()
        if len(current) >= 2 and c0 in _SECOND_HEADER_COL0 and c1 in _SECOND_HEADER_COL1:
            chunks.append(current)
            current = [row]
            continue
        current.append(row)
    chunks.append(current)
    return chunks


def is_corner_dual_header(rows: list[list[str]]) -> bool:
    """Encabezado | | ColA | ColB | (comparación Cromos/Chapería)."""
    if len(rows) < 2 or len(rows[0]) < 3:
        return False
    return rows[0][0].strip() == "" and any(c.strip() for c in rows[0][1:])


def split_corner_dual_header(rows: list[list[str]]) -> list[list[list[str]]]:
    header = rows[0]
    out: list[list[list[str]]] = []
    for ci, col_name in enumerate(header[1:]):
        name = re.sub(r"\*+", "", col_name).strip() or f"Columna {ci + 1}"
        sub: list[list[str]] = [["Aspecto", name]]
        for row in rows[1:]:
            label = row[0] if row else ""
            val = row[ci + 1] if len(row) > ci + 1 else ""
            sub.append([label, val])
        out.append(sub)
    return out


def _write_one_table(doc: Document, rows: list[list[str]]) -> None:
    # Párrafo previo no queda huérfano respecto de la tabla
    if doc.paragraphs:
        set_paragraph_keep(doc.paragraphs[-1], keep_next=True, keep_lines=True)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    for i, row in enumerate(rows):
        for j in range(cols):
            table.rows[i].cells[j].text = row[j] if j < len(row) else ""
    style_table(doc, table)
    # Párrafo Normal entre tablas: evita fusión visual
    doc.add_paragraph(style=STYLE_NORMAL)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if is_corner_dual_header(rows):
        for chunk in split_corner_dual_header(rows):
            _write_one_table(doc, chunk)
        return
    for chunk in split_dual_title_tables(rows):
        if chunk:
            _write_one_table(doc, chunk)


_LIST_RE = re.compile(r"^(?P<indent>[ \t]*)(?P<marker>[-*]|\d+\.)\s+(?P<body>.+)$")


def list_level_from_indent(indent: str) -> int:
    """0 = lista, 1 = sublista, 2 = sub-sublista (tabs o 2+ espacios)."""
    expanded = indent.expandtabs(2)
    if not expanded:
        return 0
    # 2 espacios = 1 nivel; 4 = 2
    return min(2, max(1, len(expanded) // 2))


def process_markdown(doc: Document, content: str) -> None:
    lines = content.splitlines()
    i = 0
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            if is_table_separator(stripped):
                i += 1
                continue
            table_buffer.append(parse_table_row(stripped))
            i += 1
            continue
        flush_table()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_rule_separator(doc)
            i += 1
            continue

        if stripped.startswith("#### "):
            add_body_paragraph(doc, stripped[5:].strip(), STYLE_H4)
            i += 1
            continue
        if stripped.startswith("### "):
            add_body_paragraph(doc, stripped[4:].strip(), STYLE_H3)
            i += 1
            continue
        if stripped.startswith("## "):
            add_body_paragraph(doc, stripped[3:].strip(), STYLE_H2)
            i += 1
            continue
        if stripped.startswith("# "):
            # Título 1 nativo (+ barras decorativas)
            add_part_title_ascii(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("> "):
            add_body_paragraph(doc, stripped[2:].strip(), STYLE_QUOTE)
            i += 1
            continue

        # Mecánica de profesión — celeste (#55FFFF), mismo tono que Título 3
        if stripped.startswith("**En la mesa:**") or stripped.startswith("En la mesa:"):
            add_body_paragraph(doc, stripped, STYLE_MESA)
            i += 1
            continue

        m = _LIST_RE.match(line.rstrip())
        if m and (m.group("marker") in ("-", "*") or m.group("marker").endswith(".")):
            # solo viñetas -/* como lista Word; numeradas → mismo nivel lista
            level = list_level_from_indent(m.group("indent"))
            add_bullet(doc, m.group("body").strip(), level=level)
            i += 1
            continue

        add_body_paragraph(doc, stripped, STYLE_NORMAL)
        i += 1

    flush_table()


def add_cover_page(doc: Document) -> None:
    if not PORTADA.exists():
        raise FileNotFoundError(f"Falta portada: {PORTADA}")
    art = PORTADA.read_text(encoding="utf-8").splitlines()
    while art and art[-1] == "":
        art.pop()
    add_cover_art(doc, art, "Manual de reglas — edición clara")


def main() -> None:
    doc = Document()
    setup_document_styles(doc)
    add_cover_page(doc)
    start_body_layout(doc)

    for idx, name in enumerate(CHAPTER_FILES):
        path = CAPITULOS / name
        if not path.exists():
            raise FileNotFoundError(path)
        text = path.read_text(encoding="utf-8")
        text = re.sub(r"^> \*\*Borrador.*?\n\n", "", text, flags=re.MULTILINE)
        process_markdown(doc, text)
        if idx < len(CHAPTER_FILES) - 1:
            add_rule_separator(doc)

    apply_all_section_margins(doc)
    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
