#!/usr/bin/env python3
"""Ficha canónica: ficha.docx (editada en Word). No regenerar ni recompactar.

Sincroniza docs/assets/ficha.txt + docs/ref/ficha.docx, y clona la hoja
al final del manual tal cual, sin título ni barras.
"""
from __future__ import annotations

import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCRIPTS = Path(__file__).resolve().parent
ROOT = Path(__file__).resolve().parents[2]
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from word_styles import (  # noqa: E402
    COVER_LEFT_INDENT,
    C_ACCENT,
    C_HEADING,
    C_TEXT_BRIGHT,
    FONT_COVER,
    FONT_FAMILY,
    SZ_COVER_ART,
    _apply_font,
    _preserve_run_spaces,
    _set_paragraph_shading,
    set_document_page_background,
    set_section_columns,
)

FICHA_DOCX = ROOT / "ficha.docx"
FICHA_GOLDEN = ROOT / "docs" / "ref" / "ficha.docx"
FICHA_TXT = ROOT / "docs" / "assets" / "ficha.txt"
PORTADA = ROOT / "docs" / "assets" / "portada-ascii.txt"
MANUAL = ROOT / "cyberpunk-pbta.docx"
FORMATO = ROOT / "docs" / "ref" / "formato.docx"
LAYOUT = ROOT / "docs" / "assets" / "layout.yaml"
HOJA_TITLE = "Hoja de personaje"
C_INK = RGBColor(0x00, 0x00, 0x00)
W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
_SECT_COPY = ("w:pgSz", "w:pgMar", "w:cols", "w:docGrid")
_PROMPT = "PbtA:\\>"


def export_ficha_txt(doc: Document | None = None) -> None:
    src = doc or Document(str(FICHA_DOCX))
    FICHA_TXT.write_text("\n".join(p.text for p in src.paragraphs) + "\n", encoding="utf-8")


def snapshot_ficha() -> None:
    if not FICHA_DOCX.exists():
        raise FileNotFoundError(f"Falta la ficha canónica: {FICHA_DOCX}")
    FICHA_GOLDEN.parent.mkdir(parents=True, exist_ok=True)
    if FICHA_GOLDEN.resolve() != FICHA_DOCX.resolve():
        shutil.copy2(FICHA_DOCX, FICHA_GOLDEN)


def _style_ficha_cover_para(p) -> None:
    _set_paragraph_shading(p, "FFFFFF")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pf = p.paragraph_format
    pf.left_indent = COVER_LEFT_INDENT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(8)
    pf.widow_control = False
    pPr = p._p.get_or_add_pPr()
    wrap = pPr.find(qn("w:wordWrap"))
    if wrap is None:
        wrap = OxmlElement("w:wordWrap")
        pPr.append(wrap)
    wrap.set(qn("w:val"), "0")


def _fill_cover_runs(p, line: str) -> None:
    line = line.rstrip("\n\r")
    m = re.match(r"(.*)(PbtA:\\>|pbta:\\>)(.*)$", line, flags=re.I)
    chunks: list[tuple[str, str]]
    if m:
        chunks = []
        if m.group(1):
            chunks.append((m.group(1), FONT_COVER))
        chunks.append((_PROMPT, FONT_FAMILY))
        if m.group(3):
            chunks.append((m.group(3), FONT_COVER))
    else:
        chunks = [(line, FONT_COVER)]
    for text, font in chunks:
        run = p.add_run(text)
        _apply_font(run, size=SZ_COVER_ART, bold=True, color=C_INK, font_name=font)
        _preserve_run_spaces(run)


def _logo_reversed(doc: Document) -> bool:
    if len(doc.paragraphs) < 5:
        return False
    return "PbtA:\\>" in doc.paragraphs[0].text and "_____" in doc.paragraphs[4].text


def _fix_reversed_logo(doc: Document) -> None:
    anchor = doc.paragraphs[5]._element
    els = [doc.paragraphs[i]._element for i in range(5)]
    for el in els:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    for el in reversed(els):
        anchor.addprevious(el)


def ensure_cyberpunk_header(doc: Document) -> bool:
    """Antepone el ASCII Cyberpunk + PbtA:\\> si la ficha no lo tiene."""
    if not doc.paragraphs:
        return False
    if _logo_reversed(doc):
        _fix_reversed_logo(doc)
        return True
    head = "\n".join(p.text for p in doc.paragraphs[:6])
    if "PbtA:\\>" in head or "pbta:\\>" in head.lower() or "/___/" in head:
        return False
    art = PORTADA.read_text(encoding="utf-8").splitlines()
    while art and not art[-1].strip("\xa0 "):
        art.pop()
    first = doc.paragraphs[0]._element
    created = []
    for line in art:
        p = doc.add_paragraph()
        _style_ficha_cover_para(p)
        _fill_cover_runs(p, line)
        created.append(p._element)
    for el in created:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    for el in created:
        first.addprevious(el)
    section = doc.sections[0]
    if section.top_margin is not None and int(section.top_margin) < 0:
        section.top_margin = Cm(0)
    return True


def _strip_sectpr(p_el) -> None:
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return
    sect = pPr.find(qn("w:sectPr"))
    if sect is not None:
        pPr.remove(sect)


def _blank_hf(header_or_footer) -> None:
    header_or_footer.is_linked_to_previous = False
    paras = list(header_or_footer.paragraphs)
    if not paras:
        return
    paras[0].text = ""
    for extra in paras[1:]:
        parent = extra._element.getparent()
        if parent is not None:
            parent.remove(extra._element)


def _apply_ficha_page(section, src_section) -> None:
    dst, src = section._sectPr, src_section._sectPr
    for tag in _SECT_COPY:
        old = dst.find(qn(tag))
        new = src.find(qn(tag))
        if old is not None:
            dst.remove(old)
        if new is not None:
            dst.append(deepcopy(new))
    section.different_first_page_header_footer = False
    set_section_columns(section, 1)
    _blank_hf(section.header)
    _blank_hf(section.footer)


def _set_next_page(sectPr) -> None:
    typ = sectPr.find(qn("w:type"))
    if typ is None:
        typ = OxmlElement("w:type")
        sectPr.insert(0, typ)
    typ.set(qn("w:val"), "nextPage")


def _last_body_sectpr_paragraph(doc: Document):
    for p in reversed(doc.paragraphs):
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            return p
    return None


def _ficha_block_start(doc: Document) -> int | None:
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if p.text.strip() == HOJA_TITLE:
            start = i
            if i > 0 and paras[i - 1].text.strip().startswith("====="):
                start = i - 1
            return start
    exp = next((i for i, p in enumerate(paras) if i > 200 and p.text.startswith("Experiencia:")), None)
    if exp is None:
        return next((i for i, p in enumerate(paras) if i > 200 and "┌" in p.text), None)
    j = exp
    while j > 0:
        prev = paras[j - 1]
        pPr = prev._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            break
        st = prev.style.name if prev.style else ""
        if st in ("Heading 1", "Heading 2", "Heading 3", "List Paragraph"):
            break
        j -= 1
    return j


def _already_appended(doc: Document) -> bool:
    return _ficha_block_start(doc) is not None


def _insert_ficha_clones(doc: Document, src: Document) -> None:
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        raise RuntimeError("El manual no tiene w:sectPr de sección final")
    for p in src.paragraphs:
        clone = deepcopy(p._element)
        clone.attrib.pop(f"{W14}paraId", None)
        clone.attrib.pop(f"{W14}textId", None)
        _strip_sectpr(clone)
        sectPr.addprevious(clone)


def _remove_cloned_ficha(doc: Document) -> None:
    start = _ficha_block_start(doc)
    if start is None:
        return
    for p in list(doc.paragraphs)[start:]:
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def append_hoja_de_personaje(doc: Document, ficha_path: Path | None = None) -> None:
    """Última página = clon exacto de ficha.docx, sin título."""
    src = Document(str(ficha_path or FICHA_DOCX))
    if _already_appended(doc):
        _remove_cloned_ficha(doc)
    else:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
    brk = _last_body_sectpr_paragraph(doc)
    if brk is not None:
        _set_next_page(brk._element.find(qn("w:pPr")).find(qn("w:sectPr")))
    _apply_ficha_page(doc.sections[-1], src.sections[0])
    _insert_ficha_clones(doc, src)


def extract_ficha_from_manual(manual: Document) -> Document:
    """Copia la última sección del manual a un .docx de una página."""
    start = _ficha_block_start(manual)
    if start is None:
        raise RuntimeError("No se encontró la hoja de personaje en el manual")
    out = Document()
    body = out.element.body
    for p in list(out.paragraphs):
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    _apply_ficha_page(out.sections[0], manual.sections[-1])
    set_document_page_background(out, "FFFFFF")
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        raise RuntimeError("El docx nuevo no tiene w:sectPr")
    for p in manual.paragraphs[start:]:
        clone = deepcopy(p._element)
        clone.attrib.pop(f"{W14}paraId", None)
        clone.attrib.pop(f"{W14}textId", None)
        _strip_sectpr(clone)
        sectPr.addprevious(clone)
    return out


def _rgb_hex(color: RGBColor) -> str:
    return str(color).upper()


def _write_layout_snapshot(manual: Document) -> None:
    from datetime import date

    today = date.today().isoformat()
    h = _rgb_hex(C_HEADING)
    a = _rgb_hex(C_ACCENT)
    b = _rgb_hex(C_TEXT_BRIGHT)
    lines = [
        f"# Snapshot de layout tomado de cyberpunk-pbta.docx ({today})",
        "# Referencia binaria: docs/ref/formato.docx",
        "# El builder no reproduce márgenes superiores page-a-page (ajuste manual);",
        "# sí reproduce: A4, 0,5 cm L/R/B, 2 columnas, Título 1 a 1 columna con barras =====.",
        "",
        "page: A4",
        "margins_cm:",
        "  left: 0.5",
        "  right: 0.5",
        "  bottom: 0.5",
        "  top_default: 0.5",
        "column_gap_twips: 283",
        "body_columns: 2",
        "chapter_title:",
        '  columns: 1',
        '  bar: "==================================="',
        "  bar_len: 35",
        "  heading: Heading 1",
        "  centered: true",
        "  bold: true",
        "  size_pt: 10",
        'cover_subtitle: "MANUAL DE REGLAS"',
        "",
        "headings:",
        f'  h1: {{ size_pt: 10, bold: true,  color: "{h}", keep_next: true }}',
        f'  h2: {{ size_pt: 12, bold: false, color: "{h}", keep_next: true }}',
        f'  h3: {{ size_pt: 10, bold: false, color: "{a}", keep_next: true }}',
        f'  h4: {{ size_pt: 9,  bold: false, color: "{b}", keep_next: true }}',
        f'  mesa: {{ size_pt: 8, bold: false, color: "{a}", keep_next: true }}',
        "",
        "separator_dashes: 28",
        "",
        "# Secciones del snapshot (no se reaplican al regenerar; solo documentación)",
        "snapshot_sections:",
    ]
    for s in manual.sections:
        cols_el = s._sectPr.find(qn("w:cols"))
        cols = cols_el.get(qn("w:num")) if cols_el is not None else "1"
        if cols is None:
            cols = "1"
        lines.append(f"  - {{ top_cm: {s.top_margin.cm:.2f}, cols: {cols} }}")
    LAYOUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def snapshot_from_manual() -> None:
    """Congela el Word actual (manual + hoja) como golden de futuras versiones."""
    if not MANUAL.exists():
        raise FileNotFoundError(MANUAL)
    shutil.copy2(MANUAL, FORMATO)
    manual = Document(str(MANUAL))
    ficha = extract_ficha_from_manual(manual)
    ficha.save(str(FICHA_DOCX))
    shutil.copy2(FICHA_DOCX, FICHA_GOLDEN)
    export_ficha_txt(ficha)
    _write_layout_snapshot(manual)
    print(f"OK golden {FORMATO}")
    print(f"OK {FICHA_DOCX} (extraída del manual, {len(ficha.paragraphs)} líneas)")
    print(f"OK {FICHA_GOLDEN}")
    print(f"OK {LAYOUT}")


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Snapshot de ficha / golden Word")
    parser.add_argument(
        "--from-manual",
        action="store_true",
        help="Usa cyberpunk-pbta.docx actual como golden (incluye hoja de personaje)",
    )
    args = parser.parse_args()
    if args.from_manual:
        snapshot_from_manual()
        return
    snapshot_ficha()
    export_ficha_txt()
    print(f"OK fuente {FICHA_DOCX} (sin modificar)")
    print(f"OK {FICHA_TXT}")
    print(f"OK {FICHA_GOLDEN}")


if __name__ == "__main__":
    main()
