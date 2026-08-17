#!/usr/bin/env python3
"""Exporta dumps de texto a docs/.generated/ (gitignored).

docs/ref/ es solo lectura.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
CAPITULOS = ROOT / "docs" / "capitulos"
OUT = ROOT / "docs" / ".generated"
REF = ROOT / "docs" / "ref"

CHAPTER_FILES = [
    "00-sistema.md",
    "01-crear-un-cyberpunk.md",
    "02-cyberware-reglas-y-economia.md",
    "04-catalogo-cromos.md",
    "05-catalogo-chaperia.md",
    "06-glosario.md",
]


def extract_docx(path: Path) -> str:
    doc = Document(path)
    lines: list[str] = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for ti, t in enumerate(doc.tables):
        lines.append(f"\n[TABLA {ti + 1}]")
        for row in t.rows:
            lines.append(" | ".join(c.text.replace("\n", " ").strip() for c in row.cells))
    return "\n".join(lines) + "\n"


def build_manual_completo() -> str:
    parts = [
        "pbta — MANUAL COMPLETO (texto)\n",
        "Fuente: docs/capitulos/*.md\n",
        "=" * 72 + "\n",
    ]
    for name in CHAPTER_FILES:
        path = CAPITULOS / name
        if not path.exists():
            raise FileNotFoundError(path)
        text = re.sub(r"^> \*\*Borrador.*$\n?", "", path.read_text(encoding="utf-8"), flags=re.M)
        parts.append(f"\n{'=' * 72}\nARCHIVO: {name}\n{'=' * 72}\n\n")
        parts.append(text.strip() + "\n")
    return "".join(parts)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    completo = build_manual_completo()
    (OUT / "MANUAL-completo.txt").write_text(completo, encoding="utf-8")
    print(f"OK {OUT / 'MANUAL-completo.txt'} ({len(completo)} chars)")

    orig = REF / "pbta-original.docx"
    if orig.exists():
        txt = extract_docx(orig)
        out = OUT / "pbta-original-extract.txt"
        out.write_text(txt, encoding="utf-8")
        print(f"OK {out}")


if __name__ == "__main__":
    main()
