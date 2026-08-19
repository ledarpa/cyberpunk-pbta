"""Estilos MS-DOS / pixel para cyberpunk-pbta.docx.

Usa estilos nativos de Word (Heading 1–4, Normal, List Bullet) para que
el índice automático (TOC) funcione. En Word en español se ven como
Título 1 / Título 2 / Título 3 / Normal / Lista con viñetas.
La estética VT323 + paleta DOS se aplica encima de esos estilos.
"""
from __future__ import annotations

import re
from dataclasses import dataclass

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# --- Tema (ver docs/assets/word-theme.yaml) ---
FONT_FAMILY = "VT323"                  # Terminal retro (Google Fonts, OFL)
FONT_COVER = "Courier New"             # Portada ASCII (métrica del original)
FONT_FALLBACK = "Courier New"

ASCII_WIDTH = 72
BODY_COLUMNS = 2
PAGE_W = Cm(21.0)                      # A4
PAGE_H = Cm(29.7)                      # A4
PAGE_MARGIN = Cm(0.5)                  # Márgenes exteriores
COLUMN_GAP_TWIPS = 283                 # ~0,5 cm (formato.docx)
PART_BAR_LEN = 35                      # barras ===== de Título 1 (snapshot Word)
COVER_SUBTITLE = "MANUAL DE REGLAS"
# Portada: sangría izquierda medida en pbta-original.docx (EMU 2070735)
COVER_LEFT_INDENT = Cm(5.752)
SZ_COVER_ART = Pt(7.5)                 # Courier New bold — exacto original
SZ_COVER_SUB = Pt(9)                   # Subtítulo bajo el ASCII

C_BG = RGBColor(0x0A, 0x0A, 0x0A)
C_TEXT = RGBColor(0x55, 0xFF, 0x55)
C_TEXT_BRIGHT = RGBColor(0x00, 0xFF, 0x00)
C_HEADING = RGBColor(0xFF, 0xFF, 0x55)
C_ACCENT = RGBColor(0x55, 0xFF, 0xFF)

# Escala proporcional (base cuerpo 8 pt; antes 11 pt → ratio 8/11)
SZ_BODY = Pt(8)
SZ_H1 = Pt(10)                         # Título 1 = capítulos con ===== (única negrita)
SZ_H2 = Pt(12)
SZ_H3 = Pt(10)
SZ_H4 = Pt(9)
SZ_TABLE = Pt(8)

BULLET_CHAR = "»"
BULLET_HANG = Inches(0.22)
BULLET_HANG_L2 = Inches(0.42)
TAB_COL2 = Inches(1.75)
TAB_COL3 = Inches(3.5)

# Estilos nativos Word (TOC / Título 1–3 en UI española)
STYLE_NORMAL = "Normal"
STYLE_H1 = "Heading 1"
STYLE_H2 = "Heading 2"
STYLE_H3 = "Heading 3"
STYLE_H4 = "Heading 4"
STYLE_LIST = "List Bullet"
STYLE_LIST2 = "List Bullet 2"
STYLE_QUOTE = "PBTA Quote"
STYLE_MESA = "PBTA Mesa"               # «En la mesa» — celeste como Título 3
STYLE_COVER = "PBTA Cover Line"
STYLE_PART_DECO = "PBTA Part Deco"   # barras ASCII; no entra al TOC

# numId del esquema de viñetas PBTA (abstractNum + num)
BULLET_NUM_ID = 10


@dataclass
class DocTheme:
    """Handles aplicados al construir el documento."""

    doc: Document


def _set_run_font_name(run, name: str) -> None:
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _apply_font(
    run,
    *,
    size: Pt,
    bold: bool = False,
    color: RGBColor = C_TEXT,
    font_name: str | None = None,
) -> None:
    _set_run_font_name(run, font_name or FONT_FAMILY)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color


def _preserve_run_spaces(run) -> None:
    """Evita que Word colapse espacios / NBSP al inicio o final del run."""
    t = run._element.find(qn("w:t"))
    if t is not None:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _set_paragraph_shading(paragraph, fill: str = "0A0A0A") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_outline_level(style, level: int) -> None:
    """0 = Título 1, 1 = Título 2, … — requerido para TOC automático."""
    pPr = style.element.get_or_add_pPr()
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    pPr.append(outline)


def _set_bool_ppr(pPr, tag: str, on: bool = True) -> None:
    """Activa/desactiva w:keepNext, w:keepLines, w:widowControl, etc."""
    el = pPr.find(qn(f"w:{tag}"))
    if not on:
        if el is not None:
            pPr.remove(el)
        return
    if el is None:
        el = OxmlElement(f"w:{tag}")
        pPr.append(el)
    # Word interpreta ausencia de val como true; val="true" también ok
    el.set(qn("w:val"), "true")


def set_paragraph_keep(paragraph, *, keep_next: bool = False, keep_lines: bool = False) -> None:
    """Anti-huérfano a nivel párrafo (título/bloque no se separa del siguiente)."""
    pPr = paragraph._p.get_or_add_pPr()
    if keep_next:
        _set_bool_ppr(pPr, "keepNext", True)
    if keep_lines:
        _set_bool_ppr(pPr, "keepLines", True)


def _set_style_rfonts(style, name: str) -> None:
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def apply_section_margins(section) -> None:
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN


def apply_all_section_margins(doc: Document) -> None:
    for section in doc.sections:
        apply_section_margins(section)


def set_section_columns(section, count: int) -> None:
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    if count <= 1:
        cols.set(qn("w:num"), "1")
        if qn("w:space") in cols.attrib:
            del cols.attrib[qn("w:space")]
    else:
        cols.set(qn("w:num"), str(count))
        cols.set(qn("w:space"), str(COLUMN_GAP_TWIPS))


def set_document_page_background(doc: Document, hex_color: str = "0A0A0A") -> None:
    background = doc.element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        doc.element.insert(0, background)
    background.set(qn("w:color"), hex_color)


def _ensure_paragraph_style(doc: Document, name: str) -> None:
    try:
        doc.styles[name]
    except KeyError:
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _style_paragraph(
    doc: Document,
    style_id: str,
    *,
    size: Pt,
    bold: bool = False,
    color: RGBColor = C_TEXT,
    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(3),
    left_indent=None,
    first_line_indent=None,
    tabs=None,
    outline_level: int | None = None,
    keep_next: bool = False,
    keep_lines: bool = False,
    widow_control: bool = True,
) -> None:
    st = doc.styles[style_id]
    st.font.name = FONT_FAMILY
    _set_style_rfonts(st, FONT_FAMILY)
    st.font.size = size
    st.font.bold = bold
    st.font.color.rgb = color
    pf = st.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = space_before
    pf.space_after = space_after
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if tabs:
        pf.tab_stops.clear_all()
        for pos, align in tabs:
            pf.tab_stops.add_tab_stop(pos, align)
    if outline_level is not None:
        _set_outline_level(st, outline_level)
    # Defaults anti-huérfano (ref: formato.docx — títulos pegados al bloque siguiente)
    pPr = st.element.get_or_add_pPr()
    _set_bool_ppr(pPr, "widowControl", widow_control)
    if keep_next:
        _set_bool_ppr(pPr, "keepNext", True)
    if keep_lines:
        _set_bool_ppr(pPr, "keepLines", True)


def _ensure_bullet_numbering(doc: Document) -> None:
    """Abstract numbering con » / › para lista y sublista (numId=BULLET_NUM_ID)."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering

    # Evitar duplicar si ya existe
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(BULLET_NUM_ID):
            return

    abstract_id = "10"
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)

    for ilvl, (indent_twips, hanging, glyph) in enumerate(
        (
            (440, 220, BULLET_CHAR),
            (840, 220, "›"),
            (1240, 220, "·"),
        )
    ):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), glyph)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        ind.set(qn("w:hanging"), str(hanging))
        pPr.append(ind)
        lvl.append(pPr)
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_FAMILY)
        rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
        rFonts.set(qn("w:hint"), "default")
        rPr.append(rFonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "00FF00")
        rPr.append(color)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8 pt half-points
        rPr.append(sz)
        lvl.append(rPr)
        abstract.append(lvl)

    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(BULLET_NUM_ID))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), abstract_id)
    num.append(abs_ref)
    numbering.append(num)


def _attach_num_pr(paragraph, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    # quitar numPr previo
    old = pPr.find(qn("w:numPr"))
    if old is not None:
        pPr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(BULLET_NUM_ID))
    num_pr.append(ilvl_el)
    num_pr.append(num_id)
    pPr.append(num_pr)


def setup_document_styles(doc: Document) -> DocTheme:
    """Tematiza estilos nativos Word + auxiliares PBTA; numera viñetas."""
    set_document_page_background(doc)

    for sid in (STYLE_QUOTE, STYLE_MESA, STYLE_COVER, STYLE_PART_DECO, STYLE_LIST2):
        _ensure_paragraph_style(doc, sid)

    # Normal = párrafo de cuerpo (viudas/huérfanas controladas)
    normal = doc.styles[STYLE_NORMAL]
    normal.font.name = FONT_FAMILY
    _set_style_rfonts(normal, FONT_FAMILY)
    _style_paragraph(
        doc,
        STYLE_NORMAL,
        size=SZ_BODY,
        color=C_TEXT,
        space_after=Pt(3),
        tabs=[(TAB_COL2, WD_TAB_ALIGNMENT.LEFT), (TAB_COL3, WD_TAB_ALIGNMENT.LEFT)],
        widow_control=True,
    )

    # Título 1 / 2 / 3 / 4 — keepNext+keepLines: no dejar título solo al pie de página/columna
    _style_paragraph(
        doc,
        STYLE_H1,
        size=SZ_H1,
        bold=True,
        color=C_HEADING,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        space_before=Pt(10),
        space_after=Pt(8),
        outline_level=0,
        keep_next=True,
        keep_lines=True,
    )
    _style_paragraph(
        doc,
        STYLE_H2,
        size=SZ_H2,
        bold=False,
        color=C_HEADING,
        space_before=Pt(8),
        space_after=Pt(6),
        outline_level=1,
        keep_next=True,
        keep_lines=True,
    )
    _style_paragraph(
        doc,
        STYLE_H3,
        size=SZ_H3,
        bold=False,
        color=C_ACCENT,
        space_before=Pt(6),
        space_after=Pt(4),
        outline_level=2,
        keep_next=True,
        keep_lines=True,
    )
    _style_paragraph(
        doc,
        STYLE_H4,
        size=SZ_H4,
        bold=False,
        color=C_TEXT_BRIGHT,
        space_before=Pt(4),
        space_after=Pt(0),
        outline_level=3,
        keep_next=True,
        keep_lines=True,
    )

    # Listas nativas
    for list_style, indent in ((STYLE_LIST, BULLET_HANG), (STYLE_LIST2, BULLET_HANG_L2)):
        try:
            st = doc.styles[list_style]
        except KeyError:
            _ensure_paragraph_style(doc, list_style)
            st = doc.styles[list_style]
            st.base_style = doc.styles[STYLE_LIST] if list_style == STYLE_LIST2 else doc.styles[STYLE_NORMAL]
        _style_paragraph(
            doc,
            list_style,
            size=SZ_BODY,
            color=C_TEXT,
            left_indent=indent,
            space_after=Pt(2),
            tabs=[(TAB_COL2, WD_TAB_ALIGNMENT.LEFT), (TAB_COL3, WD_TAB_ALIGNMENT.LEFT)],
        )

    _style_paragraph(
        doc,
        STYLE_QUOTE,
        size=SZ_BODY,
        color=C_ACCENT,
        left_indent=Inches(0.35),
        space_after=Pt(4),
    )
    # Líneas mecánicas de profesión — celeste; sin negrita de título
    _style_paragraph(
        doc,
        STYLE_MESA,
        size=SZ_BODY,
        bold=False,
        color=C_ACCENT,
        space_before=Pt(2),
        space_after=Pt(4),
        tabs=[(TAB_COL2, WD_TAB_ALIGNMENT.LEFT), (TAB_COL3, WD_TAB_ALIGNMENT.LEFT)],
        keep_next=True,
        keep_lines=True,
    )
    _style_paragraph(
        doc,
        STYLE_COVER,
        size=SZ_COVER_ART,
        bold=True,
        color=C_TEXT_BRIGHT,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
        space_before=Pt(0),
        space_after=Pt(0),
    )
    _style_paragraph(
        doc,
        STYLE_PART_DECO,
        size=SZ_H1,
        bold=True,
        color=C_HEADING,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        space_before=Pt(0),
        space_after=Pt(0),
        keep_next=True,
        keep_lines=True,
    )

    _ensure_bullet_numbering(doc)

    for section in doc.sections:
        apply_section_margins(section)

    return DocTheme(doc=doc)


def start_body_layout(doc: Document) -> None:
    """Tras portada: nueva sección en 2 columnas."""
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_columns(section, BODY_COLUMNS)


def begin_full_width_block(doc: Document) -> None:
    """1 columna ancho completo (títulos de parte ASCII)."""
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    apply_section_margins(section)
    set_section_columns(section, 1)


def resume_body_columns(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    apply_section_margins(section)
    set_section_columns(section, BODY_COLUMNS)


def center_ascii_line(line: str, width: int = ASCII_WIDTH) -> str:
    if len(line) >= width:
        return line
    return line.center(width)


def ascii_part_title_lines(title: str, width: int = ASCII_WIDTH) -> list[str]:
    inner = title.strip().upper()
    bar_len = max(min(len(inner) + 6, 56), 24)
    bar = "=" * bar_len
    return [center_ascii_line(bar, width), center_ascii_line(inner, width), center_ascii_line(bar, width), ""]


def add_cover_art(doc: Document, lines: list[str], subtitle: str) -> None:
    """Portada ASCII literal: sin centrar, sin rstrip; Courier New 7.5 + sangría original."""
    for _ in range(4):
        p = doc.add_paragraph()
        _set_paragraph_shading(p)
        p.paragraph_format.left_indent = COVER_LEFT_INDENT

    for raw in lines:
        line = raw.rstrip("\n\r")
        p = doc.add_paragraph(style=STYLE_COVER)
        _set_paragraph_shading(p)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = COVER_LEFT_INDENT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        _apply_font(run, size=SZ_COVER_ART, bold=True, color=C_TEXT_BRIGHT, font_name=FONT_COVER)
        _preserve_run_spaces(run)

    doc.add_paragraph()
    sub = doc.add_paragraph(style=STYLE_COVER)
    _set_paragraph_shading(sub)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.left_indent = None
    run = sub.add_run(subtitle.upper())
    _apply_font(run, size=SZ_COVER_SUB, bold=True, color=C_HEADING, font_name=FONT_FAMILY)
    doc.add_paragraph()


def add_part_title_ascii(doc: Document, title: str) -> None:
    """Capítulo (#): 1 columna, barras ===== + Título 1 centrado en negrita."""
    begin_full_width_block(doc)
    bar = "=" * PART_BAR_LEN

    p = doc.add_paragraph(style=STYLE_PART_DECO)
    _set_paragraph_shading(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(bar)
    _apply_font(run, size=SZ_H1, bold=True, color=C_HEADING)

    add_body_paragraph(doc, title.strip(), STYLE_H1, bold_parts=False)
    last = doc.paragraphs[-1]
    last.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(style=STYLE_PART_DECO)
    _set_paragraph_shading(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(bar)
    _apply_font(run, size=SZ_H1, bold=True, color=C_HEADING)

    resume_body_columns(doc)


def add_rule_separator(doc: Document) -> None:
    """Separador visual `---` — no fuerza salto de página."""
    p = doc.add_paragraph(style=STYLE_NORMAL)
    _set_paragraph_shading(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("-" * 28)
    _apply_font(run, size=SZ_BODY, color=C_ACCENT)


def add_body_paragraph(doc: Document, text: str, style: str = STYLE_NORMAL, bold_parts: bool = True) -> None:
    size_map = {
        STYLE_NORMAL: SZ_BODY,
        STYLE_H1: SZ_H1,
        STYLE_H2: SZ_H2,
        STYLE_H3: SZ_H3,
        STYLE_H4: SZ_H4,
        STYLE_QUOTE: SZ_BODY,
        STYLE_MESA: SZ_BODY,
    }
    color_map = {
        STYLE_H1: C_HEADING,
        STYLE_H2: C_HEADING,
        STYLE_H3: C_ACCENT,
        STYLE_H4: C_TEXT_BRIGHT,
        STYLE_QUOTE: C_ACCENT,
        STYLE_MESA: C_ACCENT,
    }
    size = size_map.get(style, SZ_BODY)
    default_color = color_map.get(style, C_TEXT)
    # Solo Título 1 (bloques =====) lleva negrita de título
    heading_bold = style == STYLE_H1

    p = doc.add_paragraph(style=style)
    _set_paragraph_shading(p)
    if bold_parts:
        parts = __import__("re").split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                # En mesa: negrita del mismo celeste; en cuerpo: verde brillante
                emph = default_color if style == STYLE_MESA else (
                    C_TEXT_BRIGHT if not heading_bold else default_color
                )
                _apply_font(run, size=size, bold=True, color=emph)
            elif part:
                run = p.add_run(part)
                _apply_font(run, size=size, bold=heading_bold, color=default_color)
    else:
        run = p.add_run(text)
        _apply_font(run, size=size, bold=heading_bold, color=default_color)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Lista (level 0) o sublista (level 1+); estilo List Bullet + numPr nativo."""
    level = max(0, min(level, 2))
    style = STYLE_LIST2 if level >= 1 else STYLE_LIST
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style=STYLE_LIST)
    _set_paragraph_shading(p)
    _attach_num_pr(p, level)

    parts = __import__("re").split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _apply_font(run, size=SZ_BODY, bold=True, color=C_TEXT_BRIGHT)
        elif part:
            run = p.add_run(part)
            _apply_font(run, size=SZ_BODY, color=C_TEXT)


def style_table(doc: Document, table) -> None:
    """Tabla con estilo Table Grid + cabecera + filas que no se parten (anti-huérfano)."""
    table.style = "Table Grid"
    if table.rows:
        tr = table.rows[0]._tr
        tr_pr = tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            hdr = OxmlElement("w:tblHeader")
            tr_pr.append(hdr)

    for i, row in enumerate(table.rows):
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        # No partir fila entre páginas/columnas
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            _set_cell_shading(cell, "0A0A0A" if i else "00AA00")
            for p in cell.paragraphs:
                try:
                    p.style = doc.styles[STYLE_NORMAL]
                except KeyError:
                    pass
                _set_paragraph_shading(p, "0A0A0A" if i else "00AA00")
                for run in p.runs:
                    _apply_font(
                        run,
                        size=SZ_TABLE,
                        bold=(i == 0),
                        color=RGBColor(0, 0, 0) if i == 0 else C_TEXT,
                    )


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _brand_pbta(text: str) -> str:
    """pbta visible → PbtA. No toca nombres de archivo cyberpunk-pbta.*"""
    if not text or "pbta" not in text.lower():
        return text
    parts = re.split(r"(cyberpunk-pbta\S*)", text, flags=re.I)
    out = []
    for i, part in enumerate(parts):
        if i % 2 == 1:
            out.append(part)
        else:
            out.append(re.sub(r"pbta", "PbtA", part, flags=re.I))
    return "".join(out)


def _iter_story_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from hf.paragraphs


def replace_pbta_branding(doc: Document) -> int:
    """Sustituye pbta → PbtA en cuerpo, tablas, encabezados y pies."""
    n = 0
    for p in _iter_story_paragraphs(doc):
        for run in p.runs:
            old = run.text
            new = _brand_pbta(old)
            if new != old:
                run.text = new
                _preserve_run_spaces(run)
                n += 1
    return n
